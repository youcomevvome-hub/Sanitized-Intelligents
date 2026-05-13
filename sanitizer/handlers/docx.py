"""DOCX handler.

Strategy:
  1. Walk paragraphs + tables, run PII analyzer on the text, replace matched
     spans with [REDACTED] tokens. This handles names, emails, IDs, numbers,
     custom regex and keywords.
  2. For every embedded image inside the .docx (`word/media/*`), run the
     full image sanitizer in-place and rewrite the zip archive.
"""
from __future__ import annotations

import io
import re
import shutil
import zipfile
from pathlib import Path
from typing import List

import cv2
import numpy as np

from sanitizer.core.types import Detection, DetectionKind, SanitizeRequest, SanitizeResult
from sanitizer.detectors.base import FaceDetector, TextDetector
from sanitizer.detectors.pii_text import PIITextAnalyzer
from sanitizer.handlers.image import sanitize_image
from sanitizer.utils import logger


_NUMBER_RE = re.compile(r"\d{2,}")
REDACTION_TOKEN = "[REDACTED]"


def _apply_custom_replacements(text: str, replacements: dict[str, str]) -> str:
    out = text
    for src, dst in (replacements or {}).items():
        if not src:
            continue
        out = re.sub(re.escape(src), dst or "", out, flags=re.IGNORECASE)
    return out


def _redact_text(text: str, pii: PIITextAnalyzer, request: SanitizeRequest) -> tuple[str, List[Detection]]:
    if not text:
        return text, []
    text = _apply_custom_replacements(text, request.custom_replacements)
    replacement_token = (request.replacement_text or REDACTION_TOKEN).strip() or REDACTION_TOKEN
    detections: List[Detection] = []
    # Collect spans
    spans: List[tuple[int, int, str]] = []

    if request.redact_pii:
        for s in pii.analyze(text, request.custom_patterns, request.keywords):
            spans.append((s.start, s.end, s.label))
    if request.redact_numbers:
        for m in _NUMBER_RE.finditer(text):
            spans.append((m.start(), m.end(), "NUMBER"))

    if not spans:
        return text, []

    # Merge overlaps
    spans.sort()
    merged: List[tuple[int, int, str]] = []
    for s, e, lbl in spans:
        if merged and s <= merged[-1][1]:
            ps, pe, plbl = merged[-1]
            merged[-1] = (ps, max(pe, e), plbl)
        else:
            merged.append((s, e, lbl))

    # Rebuild text + record detections (bbox = char offsets)
    out_parts: List[str] = []
    cursor = 0
    for s, e, lbl in merged:
        out_parts.append(text[cursor:s])
        out_parts.append(replacement_token)
        detections.append(
            Detection(
                kind=DetectionKind.PII if lbl != "NUMBER" else DetectionKind.NUMBER,
                bbox=(s, e, 0, 0),
                label=lbl,
            )
        )
        cursor = e
    out_parts.append(text[cursor:])
    return "".join(out_parts), detections


def _redact_doc_paragraphs(doc, pii: PIITextAnalyzer, request: SanitizeRequest) -> List[Detection]:
    detections: List[Detection] = []

    def _process_paragraph(p) -> None:
        full_text = "".join(run.text or "" for run in p.runs)
        new_text, dets = _redact_text(full_text, pii, request)
        if new_text != full_text:
            # Wipe existing runs and put the redacted text into the first run
            if p.runs:
                p.runs[0].text = new_text
                for run in p.runs[1:]:
                    run.text = ""
            detections.extend(dets)

    for p in doc.paragraphs:
        _process_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _process_paragraph(p)
    return detections


def _sanitize_embedded_images(
    docx_path: Path,
    out_path: Path,
    face_det: FaceDetector,
    text_det: TextDetector,
    pii: PIITextAnalyzer,
    request: SanitizeRequest,
) -> List[Detection]:
    """Rewrite the docx zip, replacing each image inside `word/media/`
    with its sanitized version."""
    tmp_dir = out_path.parent / f".sanitize_tmp_{out_path.stem}"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    image_dets: List[Detection] = []

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/media/") and item.filename.lower().endswith(
                    (".png", ".jpg", ".jpeg", ".bmp", ".tiff")
                ):
                    arr = np.frombuffer(data, dtype=np.uint8)
                    img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
                    if img is not None:
                        # Write to temp, sanitize, read back, replace bytes
                        in_tmp = tmp_dir / Path(item.filename).name
                        out_tmp = tmp_dir / f"out_{Path(item.filename).name}"
                        cv2.imwrite(str(in_tmp), img)
                        try:
                            res = sanitize_image(in_tmp, out_tmp, face_det, text_det, pii, request)
                            image_dets.extend(res.detections)
                            data = out_tmp.read_bytes()
                        except Exception as e:
                            logger.warning(f"Failed to sanitize embedded image {item.filename}: {e}")
                zout.writestr(item, data)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return image_dets


def sanitize_docx(
    input_path: str | Path,
    output_path: str | Path,
    face_det: FaceDetector,
    text_det: TextDetector,
    pii: PIITextAnalyzer,
    request: SanitizeRequest,
) -> SanitizeResult:
    import docx as python_docx

    input_path = Path(input_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # 1. Redact text — save to an intermediate file
    intermediate = output_path.with_suffix(".text-redacted.docx")
    doc = python_docx.Document(str(input_path))
    text_dets = _redact_doc_paragraphs(doc, pii, request)
    doc.save(str(intermediate))

    # 2. Sanitize embedded images by rewriting the zip
    image_dets = _sanitize_embedded_images(intermediate, output_path, face_det, text_det, pii, request)
    intermediate.unlink(missing_ok=True)

    return SanitizeResult(
        output_path=str(output_path),
        detections=text_dets + image_dets,
        media_type="docx",
        meta={"text_redactions": len(text_dets), "image_redactions": len(image_dets)},
    )
